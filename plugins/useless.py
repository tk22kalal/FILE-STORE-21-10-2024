from bot import Bot
import openai
from pyrogram.types import Message
from pyrogram import filters
from pyrogram import Client
import io
from google.cloud import vision
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
import pdf2image
import asyncio
import re
from groq import Groq

# Configure Google Gemini API and Vision
openai.api_key = "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"
genai.configure(api_key="AIzaSyCL_5XEd39cgAdcIBLhbu9OaT-RrhSSSjI")
vision_client = vision.ImageAnnotatorClient.from_service_account_file("plugins/gen-lang-client-0707503202-21d07fd84f57.json")

@Client.on_message(filters.document)
async def pdf_handler(client: Client, message: Message):
    """Handle PDF uploads, extract text via OCR, convert to notes format, and send as a formatted Word document."""
    if message.document.file_name.endswith(".pdf"):
        # Send a temporary "Processing..." message
        processing_message = await client.send_message(chat_id=message.chat.id, text="Processing... 0%")

        file_id = message.document.file_id
        file = await client.download_media(file_id)
        
        document = Document()

        try:
            # Convert each page to an image and perform OCR
            images = pdf2image.convert_from_path(file, dpi=300)
            total_pages = len(images)

            for page_num, image in enumerate(images, start=1):
                # Update the progress percentage
                progress = int((page_num / total_pages) * 100)
                await processing_message.edit_text(f"Processing... {progress}%")

                # Process the current page
                image_bytes = io.BytesIO()
                image.save(image_bytes, format='JPEG')
                vision_image = vision.Image(content=image_bytes.getvalue())
                
                ocr_response = vision_client.text_detection(image=vision_image)
                ocr_text = ocr_response.full_text_annotation.text
                
                # Generate notes using Gemini API
                formatted_prompt = (
                    "Analyze the following content and perform the following tasks step by step:\n\n"
                    "1. Identify and extract the **Main Title**, **Headings**, **Sub-headings**, and the structure of the content. Clearly mark each section using:\n"
                    "   - Use '###' for Main Title.\n"
                    "   - Use '***' for Headings.\n"
                    "   - Use '##' for Sub-headings.\n"
                    "   - Use '-' for normal paragraph key points.\n\n"
                    "2. Simplify and explain the content into **point-wise format** under each identified section, while maintaining the **original meaning**.\n"
                    "   - Organize the points clearly and concisely.\n"
                    "   - Break down complex information into smaller, understandable parts.\n\n"
                    "3. Ensure all **formulas**, **tables**, **cycles**, and **mind maps** are preserved **as they are** without modification or simplification.\n\n"
                    "4. Add bullet points for each section with the following rules:\n"
                    "   - Use numbered bullets for Headings.\n"
                    "   - Use '☆' for Sub-headings.\n"
                    "   - Use '●' for Points.\n"
                    "   - Use '⭘' for Key Points.\n"
                    "   - Use '◊' for Sub-Points.\n\n"
                    "5. Highlight important words or phrases in each line using **bold formatting**.\n"
                    "   - Ensure that highlighted words add clarity or emphasize critical information.\n\n"
                    "Finally, ensure the output is cleanly structured and well-formatted while preserving the meaning of the original content:\n\n"
                    + ocr_text
                )
                client = Groq()

                # Generate the completion
                completion = client.chat.completions.create(
                    model="llama-3.1-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are an expert assistant that structures and simplifies content."},
                        {"role": "user", "content": formatted_prompt}
                    ]
                )
                
                # Extract the generated text
                notes_text = completion["choices"][0]["message"]["content"]

                # Add the notes to the Word document with formatting
                lines = notes_text.split("\n")
                for line in lines:
                    if line.startswith("###"):
                        # Main Heading formatting with zero left indent
                        heading = document.add_paragraph(style='Title')
                        heading_format = heading.paragraph_format
                        heading_format.left_indent = Pt(0)
                        heading_format.space_before = Pt(0)
                        heading_format.space_after = Pt(0)
                        run = heading.add_run(line.replace("###", "").strip())
                        run.font.bold = True
                        run.font.size = Pt(28)
                        run.font.name = 'Baskerville Old Face'
                        run.font.color.rgb = RGBColor(255, 165, 0)  # Green

                    elif line.startswith("***"):
                        # H2 Heading formatting with slight indent
                        heading = document.add_paragraph(style='Heading 1')
                        heading_format = heading.paragraph_format
                        heading_format.left_indent = Pt(0)  # Slight indent
                        heading_format.space_before = Pt(0)
                        heading_format.space_after = Pt(0.5)
                        run = heading.add_run(line.replace("***", "").strip())
                        run.font.bold = True
                        run.font.size = Pt(15)
                        run.font.name = 'Tahoma'
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green

                    elif line.startswith("##"):
                        # H3 Subheading formatting with moderate indent
                        heading = document.add_paragraph(style='Heading 2')
                        heading_format = heading.paragraph_format
                        heading_format.left_indent = Pt(0)  # Moderate indent
                        heading_format.space_before = Pt(0)
                        heading_format.space_after = Pt(0.5)
                        run = heading.add_run(line.replace("##", "").strip())
                        run.font.bold = True
                        run.font.size = Pt(15)
                        run.font.name = 'Tahoma'
                        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange

                    elif line.startswith("-"):
                        # Normal Paragraph Text
                        paragraph = document.add_paragraph(style='List Bullet')
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Pt(15)
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(1)
                        run = paragraph.add_run(line.replace("-", "").strip())
                        run.font.size = Pt(13)
                        run.font.name = 'Tahoma'

                    else:
                        paragraph = document.add_paragraph()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Pt(15)  # Further indent for non-marked text
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0.5)
                        run = paragraph.add_run(line.strip())
                        run.font.size = Pt(13)
                        run.font.name = 'Tahoma'

                # Rest between pages to avoid hitting API limits or overloading
                await asyncio.sleep(2)

            # Post-process document for **bold** text formatting
            for paragraph in document.paragraphs:
                # Find all occurrences of **bold text**
                parts = re.split(r'(\*\*[^*]+\*\*)', paragraph.text)
                paragraph.clear()
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        bold_text = part.replace("**", "")
                        run = paragraph.add_run(bold_text)
                        run.font.bold = True  # Strong bold
                    else:
                        run = paragraph.add_run(part)
                    run.font.size = Pt(13)
                    run.font.name = 'Tahoma'
                    
            # Save the Word document
            word_file = io.BytesIO()
            document.save(word_file)
            word_file.seek(0)

            # Send the Word document back to the user
            await client.send_document(
                chat_id=message.chat.id,
                document=word_file,
                file_name="Formatted_Notes.docx",
                caption="Here are your notes in a structured Microsoft Word format."
            )

            # Delete the temporary processing message
            await processing_message.delete()

        except Exception as e:
            await message.reply("Error processing the PDF. Please try again.")
            print(f"Error processing PDF: {e}")
            await processing_message.delete()

    else:
        await message.reply("Please upload a valid PDF document.")
